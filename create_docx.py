# -*- coding: utf-8 -*-
import os
import re
import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 콘솔 한글 인코딩 대응
sys.stdout.reconfigure(encoding='utf-8')

# XML 조작용 헬퍼 함수
def set_cell_background(cell, fill_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180): # dxa 단위 (1 pt = 20 dxa)
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table):
    tblPr = table._tbl.tblPr
    borders_xml = (
        '<w:tblBorders ' + nsdecls('w') + '>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="D1D5DB"/>'
        '<w:bottom w:val="single" w:sz="6" w:space="0" w:color="9CA3AF"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="E5E7EB"/>'
        '<w:insideV w:val="none"/>'
        '<w:left w:val="none"/>'
        '<w:right w:val="none"/>'
        '</w:tblBorders>'
    )
    tblPr.append(parse_xml(borders_xml))

# 단락 왼쪽 테두리(데코레이션용) 설정
def set_paragraph_left_border(paragraph, hex_color="0D9488", size_eighth_pt=24):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), str(size_eighth_pt))
    left.set(qn('w:space'), '12')
    left.set(qn('w:color'), hex_color)
    pBdr.append(left)

# 단락 하단 테두리(수평선) 설정
def set_paragraph_bottom_border(paragraph, hex_color="CBD5E1", size_eighth_pt=6):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size_eighth_pt))
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), hex_color)
    pBdr.append(bottom)

# 단락 배경색(음영) 설정
def set_paragraph_shading(paragraph, fill_hex):
    pPr = paragraph._p.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    pPr.append(shd)

# 폰트 스타일 설정 (한글 맑은 고딕 / 영문 혼용 대응)
def set_font_style(run, name='맑은 고딕', size_pt=10.5, bold=False, italic=False, rgb_color=None):
    run.font.name = name
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if rgb_color:
        run.font.color.rgb = rgb_color

# 마크다운 인라인 요소 파싱 및 런 추가
def add_markdown_runs(paragraph, text, size_pt=10.5, bold=False, italic=False, rgb_color=None):
    # **볼드** 및 `코드` 파싱
    parts = re.split(r'(\*\*.*?\*\*|`.*?`)', text)
    for part in parts:
        if not part:
            continue
        
        current_bold = bold
        current_italic = italic
        current_color = rgb_color
        current_font = '맑은 고딕'
        current_size = size_pt
        
        if part.startswith('**') and part.endswith('**'):
            part_text = part[2:-2]
            current_bold = True
        elif part.startswith('`') and part.endswith('`'):
            part_text = part[1:-1]
            current_font = 'Consolas'
            current_size = size_pt - 0.5
            # 코드조각용 Rose 테마 색상 부여
            current_color = RGBColor(190, 24, 74) 
        else:
            part_text = part
            
        run = paragraph.add_run(part_text)
        set_font_style(run, name=current_font, size_pt=current_size, 
                       bold=current_bold, italic=current_italic, rgb_color=current_color)

def parse_markdown_to_docx(md_path, docx_path):
    doc = Document()
    
    # 여백 설정 (상하좌우 1인치)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
    
    # 기본 Normal 스타일 폰트 초기화
    normal_style = doc.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = '맑은 고딕'
    normal_font.size = Pt(10.5)
    normal_font.color.rgb = RGBColor(30, 41, 59) # Slate 800
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n')
        stripped = line.strip()
        
        # 1. 코드 블록 멀티라인 처리
        if stripped.startswith('```'):
            code_lang = stripped[3:].strip()
            code_buffer = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_buffer.append(lines[i].rstrip('\n'))
                i += 1
            if i < len(lines):
                i += 1 # 닫는 ``` 소비
            render_code_block(doc, code_buffer, code_lang)
            continue
            
        # 2. 테이블 멀티라인 처리
        if stripped.startswith('|'):
            table_rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_rows.append(lines[i].strip())
                i += 1
            render_table(doc, table_rows)
            continue
            
        # 3. 블록인용구 멀티라인 처리
        if stripped.startswith('>'):
            quote_buffer = []
            while i < len(lines) and (lines[i].strip().startswith('>') or not lines[i].strip()):
                if lines[i].strip().startswith('>'):
                    # '>' 기호 뒤 첫 공백 제거
                    content = lines[i].strip()[1:]
                    if content.startswith(' '):
                        content = content[1:]
                    quote_buffer.append(content)
                else:
                    quote_buffer.append("")
                i += 1
            render_blockquote(doc, quote_buffer)
            continue
            
        # 4. 빈 줄 처리
        if not stripped:
            i += 1
            continue
            
        # 5. 제목 (Headings) 처리
        if stripped.startswith('# '):
            text = stripped[2:]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.keep_with_next = True
            
            # 대제목은 조금 더 크고 두껍게
            run = p.add_run(text)
            set_font_style(run, name='맑은 고딕', size_pt=22, bold=True, rgb_color=RGBColor(15, 23, 42))
            # 굵고 어두운 하단 구분 데코 선 추가
            set_paragraph_bottom_border(p, hex_color="0F172A", size_eighth_pt=12)
            
        elif stripped.startswith('## '):
            text = stripped[3:]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            
            # 중제목: 포인트 컬러인 Teal 600 적용
            add_markdown_runs(p, text, size_pt=14, bold=True, rgb_color=RGBColor(13, 148, 136))
            # 얇은 하단 구분선
            set_paragraph_bottom_border(p, hex_color="E2E8F0", size_eighth_pt=4)
            
        elif stripped.startswith('### '):
            text = stripped[4:]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            
            add_markdown_runs(p, text, size_pt=11.5, bold=True, rgb_color=RGBColor(30, 41, 59))
            
        elif stripped.startswith('#### '):
            text = stripped[5:]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            
            add_markdown_runs(p, text, size_pt=10.5, bold=True, rgb_color=RGBColor(71, 85, 105))
            
        # 6. 마크다운 수평선 처리
        elif stripped == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            set_paragraph_bottom_border(p, hex_color="CBD5E1", size_eighth_pt=6)
            
        # 7. 리스트 처리
        elif stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:]
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.2
            add_markdown_runs(p, text)
            
        # 8. 일반 본문 단락 처리
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.25
            add_markdown_runs(p, stripped)
            
        i += 1
        
    doc.save(docx_path)
    print(f"DOCX successfully generated at: {docx_path}")

def render_blockquote(doc, quote_lines):
    full_text = "\n".join(quote_lines).strip()
    if not full_text:
        return
        
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.2
    
    # 테일 테마 스타일링: 두꺼운 왼쪽 테두리와 밝은 틴트 배경
    set_paragraph_left_border(p, hex_color="0D9488", size_eighth_pt=24) # 3pt 두께
    set_paragraph_shading(p, "F0FDFA") # Teal 50
    
    # 텍스트 추가 (이탤릭, Slate 600)
    add_markdown_runs(p, full_text, size_pt=10, italic=True, rgb_color=RGBColor(71, 85, 105))

def render_code_block(doc, code_lines, code_lang):
    full_code = "\n".join(code_lines)
    if not full_code.strip():
        return
        
    # 1행 1열 테이블로 둥근 회색 컨테이너 효과 모사
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.0)
    
    # 연회색 배경 및 여백
    set_cell_background(cell, "F8FAFC") # Slate 50
    set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
    
    # 얇은 외곽 테두리
    tcPr = cell._tc.get_or_add_tcPr()
    borders_xml = (
        '<w:tcBorders ' + nsdecls('w') + '>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
        '</w:tcBorders>'
    )
    tcPr.append(parse_xml(borders_xml))
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.05
    
    display_lang = code_lang.upper() if code_lang else "CODE"
    # 만약 Mermaid 블록이면 식별 헤더 표시
    if "MERMAID" in display_lang:
        run_title = p.add_run(f"[{display_lang} STRUCTURE]\n")
        set_font_style(run_title, name='Consolas', size_pt=8.5, bold=True, rgb_color=RGBColor(100, 116, 139))
        
    run = p.add_run(full_code)
    set_font_style(run, name='Consolas', size_pt=9, rgb_color=RGBColor(51, 65, 85))

def render_table(doc, rows_data):
    cleaned_rows = []
    for r in rows_data:
        if re.search(r'^\|\s*[:-]+\s*\|', r) or '---' in r:
            continue
        cleaned_rows.append(r)
        
    if not cleaned_rows:
        return
        
    first_row_cols = [c.strip() for c in cleaned_rows[0].split('|')[1:-1]]
    num_cols = len(first_row_cols)
    if num_cols == 0:
        return
        
    table = doc.add_table(rows=len(cleaned_rows), cols=num_cols)
    set_table_borders(table)
    table.autofit = False
    
    # 열 너비 비율 최적화 (2열 및 3열)
    if num_cols == 2:
        col_widths = [Inches(1.5), Inches(4.5)]
    elif num_cols == 3:
        col_widths = [Inches(1.2), Inches(2.4), Inches(2.4)]
    else:
        col_widths = [Inches(6.0 / num_cols)] * num_cols
        
    for r_idx, row_str in enumerate(cleaned_rows):
        columns = [c.strip() for c in row_str.split('|')[1:-1]]
        while len(columns) < num_cols:
            columns.append("")
            
        for c_idx in range(num_cols):
            cell = table.cell(r_idx, c_idx)
            cell.width = col_widths[c_idx]
            
            set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
            
            cell_text = columns[c_idx]
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            
            if r_idx == 0:
                # 헤더 스타일: 다크 슬레이트 900
                set_cell_background(cell, "0F172A")
                run = p.add_run(cell_text)
                set_font_style(run, name='맑은 고딕', size_pt=9.5, bold=True, rgb_color=RGBColor(255, 255, 255))
            else:
                # 얼룩말 패턴
                if r_idx % 2 == 1:
                    set_cell_background(cell, "F8FAFC")
                else:
                    set_cell_background(cell, "FFFFFF")
                    
                add_markdown_runs(p, cell_text, size_pt=9.5)

if __name__ == "__main__":
    md_file = r"C:\Users\jang9\OneDrive\Desktop\작업\3차 프로젝트\00.메인_기획서\3차_프로젝트_기획_제안서_Notion.md"
    docx_file = r"C:\Users\jang9\OneDrive\Desktop\작업\3차 프로젝트\00.메인_기획서\3차_프로젝트_기획_제안서.docx"
    
    if os.path.exists(md_file):
        parse_markdown_to_docx(md_file, docx_file)
    else:
        print(f"Error: Markdown file not found at {md_file}")
